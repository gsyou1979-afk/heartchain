# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page margins ───
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(3.0)

# ─── Styles helpers ───
def set_run_font(run, name='맑은 고딕', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level=1, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 11}
    set_run_font(run, size=sizes.get(level, 11), bold=True)
    return p

def body(doc, text, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run)
    return p

def bi(doc, kr, cn, indent=0):
    """Bilingual paragraph: Korean + Chinese side by side (tab-separated)"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r1 = p.add_run(kr)
    r1.bold = True
    set_run_font(r1, size=10)
    p.add_run('  /  ')
    r2 = p.add_run(cn)
    set_run_font(r2, size=10, color=(100, 100, 100))
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # Rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def add_shaded_table_row(table, row_idx, color_hex='D9E1F2'):
    """Apply background shading to a row"""
    row = table.rows[row_idx]
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

# ═══════════════════════════════════════════════════
#  TITLE
# ═══════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('하트체인(HeartChain) 파트너십 계약서')
r.bold = True; r.font.size = Pt(18)
r.font.name = '맑은 고딕'

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run('心链（HeartChain）合作伙伴协议')
r2.bold = True; r2.font.size = Pt(14)
r2.font.name = '맑은 고딕'

doc.add_paragraph()

# ─── PARTIES ───
heading(doc, '계약당사자 / 合同双方', level=1, center=True)
doc.add_paragraph()

def party_block(doc, party, name_kr, name_cn, addr_kr, addr_cn, rep_kr, rep_cn, reg_kr, reg_cn):
    p = doc.add_paragraph()
    r = p.add_run(f'【{party} ({"甲方" if party=="갑" else "乙方"})】')
    r.bold = True; r.font.size = Pt(11)
    body(doc, f'상호/名称: {name_kr} ({name_cn})')
    body(doc, f'주소/地址: {addr_kr} / {addr_cn}')
    body(doc, f'대표자/法定代表人: {rep_kr} / {rep_cn}')
    body(doc, f'사업자번호/营业执照号: {reg_kr} / {reg_cn}')
    doc.add_paragraph()

party_block(doc, '갑', '____________', '____________', '____________', '____________', '____________', '____________', '____________', '____________')
party_block(doc, '을', '____________', '____________', '____________', '____________', '____________', '____________', '____________', '____________')

# ═══════════════════════════════════════════════════
#  PREAMBLE
# ═══════════════════════════════════════════════════
heading(doc, '서문 / 前言', level=1)
body(doc, '「갑」은 블록체인 기반 자원봉사 관리 플랫폼「HeartChain(하트체인/心链)」을 개발·운영하며,「을」는 __________(이하「자원봉사 단체」)를 운영하는 파트너로서 본 계약에 참여한다. 양측은 상호 이해득실과 공동 번영을 도모하며, 다음과 같이 계약을 체결한다.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 1 – Definitions
# ═══════════════════════════════════════════════════
heading(doc, '제1조 (정 의) / 第1条（定义）', level=1)
add_table(doc,
    ['용어 / 术语', '정의 / 定义'],
    [
        ['플랫폼', '「갑」이 운영하는 HeartChain 온라인 서비스 및 관련 APP'],
        ['HRT', '플랫폼 내부 토큰「Heart Token」. 현 시점 단가는 미부여, 상장 시점 역산에 따름'],
        ['유효회원', '플랫폼에 가입 후 6개월 이내 최소 1회 APP에 로그인한 회원'],
        ['단체 등록인원', '「을」를 통해 플랫폼에 가입된 유효회원 수'],
        ['개인이머니 (人头费)', '개인이 자체 추천·유치하여 가입시킨 유효회원 수'],
        ['임무 / 任务', '플랫폼에 게시된 자원봉사 활동 정보'],
        ['임무가치 / 任务价值', '임무 게시 시 발행자가 설정한 HRT 표기 수치'],
    ]
)
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 2 – Distribution Ratios
# ═══════════════════════════════════════════════════
heading(doc, '제2조 (HRT 발행总量 및 분배 비율) / 第2条（HRT发行总量及分配比例）', level=1)
body(doc, '「갑」은 플랫폼 토탈 HRT 중 아래와 같이 분배한다.')
doc.add_paragraph()

dist = doc.add_table(rows=4, cols=2)
dist.style = 'Table Grid'
dist.alignment = WD_TABLE_ALIGNMENT.CENTER
dist_data = [
    ('HRT 총 발행량', '100%'),
    ('├─ 30% → 개발자 보유 (Developer Reserve)', '300,000 HRT (예시)'),
    ('├─ 50% → 단체 풀 (Organization Pool)', '「을」 등 파트너 단체에人头方式分配'),
    ('└─ 20% → 개인 풀 (Individual Pool)', '가입시킨 개인회원에게人头方式分配'),
]
for i, (label, val) in enumerate(dist_data):
    dist.rows[i].cells[0].text = label
    dist.rows[i].cells[1].text = val
    for c in dist.rows[i].cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
if len(dist.rows) > 0:
    add_shaded_table_row(dist, 0, 'D9E1F2')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 3 – Organization Pool
# ═══════════════════════════════════════════════════
heading(doc, '제3조 (단체 풀 HRT 부여 방식) / 第3条（团体池HRT赋予方式）', level=1)

heading(doc, '3.1 기본 원칙 / 基本原则', level=2)
body(doc, '「을」에게 부여되는 단체 풀 HRT는 가입 시 1회성으로 지급되며, 다음 산식에 따른다.')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('단체 HRT 지급总量 = 해당 단체 注册人头数 × 1인당 HRT 단가')
r.bold = True; r.font.size = Pt(11)
r.font.name = 'Consolas'
doc.add_paragraph()

heading(doc, '3.2 1인당 HRT 단가 산정 / 人均 HRT 单价确定', level=2)
body(doc, '1인당 HRT 단가는 상장 시점 역산 방식으로 결정한다.')
add_table(doc,
    ['산식 요소', '설명'],
    [
        ['분자', 'HRT 총 발행량 × 50% = 단체 풀总量'],
        ['분모', '계약 체결 기준 모든 파트너 단체의 유효회원 합계'],
        ['단가 산출', '분자 ÷ 분모 = 1인당 HRT 단가'],
        ['상장 시 재검증', '「갑」은 토큰上市 시 시장 환산가치에 기반하여 최종 단가를 확정하고「을」에게 서면으로 통보'],
    ]
)
doc.add_paragraph()

heading(doc, '3.3 지급 시기 및 방법 / 发放时间及方式', level=2)
body(doc, '① 「을」는 회원 가입 정보를 플랫폼에 입력한다.\n② 「갑」은 해당 정보를 월 1회 검증·확인한다.\n③ 유효회원으로 판정된人数에 대해 해당월 말일 기준 30일 이내，「을」의 HRT 지갑으로 전송한다.\n④ 지급 완료 시「갑」은「을」에게 지급내역 명세서(人头数·단가·총액)를发送电子邮件/书面通报한다.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 4 – Individual Pool
# ═══════════════════════════════════════════════════
heading(doc, '제4조 (개인 풀 HRT 부여 방식) / 第4条（个人池HRT赋予方式）', level=1)
body(doc, '개인이 자체 추천·유치하여 플랫폼에 가입시킨 유효회원이 있는 개인회원에게 개인 풀(20%) 내에서 HRT를 부여한다. 산식 및 단가는 제3조와 동일하게 적용한다.')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('개인 HRT = 해당 개인이 유치한 유효회원수 × 1인당 HRT 단가')
r.bold = True; r.font.size = Pt(11)
r.font.name = 'Consolas'
doc.add_paragraph()
body(doc, '개인 회원에게直接 지급하며, 지급내역은 개인 지갑 주소로 자동 반영된다.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 5 – Valid Member Definition
# ═══════════════════════════════════════════════════
heading(doc, '제5조 (유효회원 인정 기준) / 第5条（有效会员认定标准）', level=1)

heading(doc, '5.1 유효회원의 정의 / 有效会员定义', level=2)
body(doc, '다음 조건을 모두 충족한 회원만 유효회원으로 인정된다.')
body(doc, '① 플랫폼에 정식 가입 완료한 자', indent=0.5)
body(doc, '② 가입 후 6개월 이내 최소 1회 이상 APP에 로그인한 자', indent=0.5)
body(doc, '③ 플랫폼 이용약관 및 관련 규정을 위반하지 않은 자', indent=0.5)
doc.add_paragraph()

heading(doc, '5.2 검증 주기 / 验证周期', level=2)
body(doc, '「갑」은 매월 1일 기준으로 유효회원 수를 점검하고, 해당 수치를「을」에게 통보한다.')
doc.add_paragraph()

heading(doc, '5.3 유효회원에 포함되지 않는 경우 / 不计入有效会员的情况', level=2)
body(doc, '① 가입 후 6개월 내 APP 미접속\n② 개인 정보가 위조·탈퇴된 것으로 확인된 경우\n③ 플랫폼 정책 위반으로 정지·삭제된 계정')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 6 – Task Incentives
# ═══════════════════════════════════════════════════
heading(doc, '제6조 (임무 관련 장려금) / 第6条（任务相关奖励金）', level=1)

heading(doc, '6.1 회원 임무완성 장려금 (团体奖励 / 任务完成奖励)', level=2)
body(doc, '「을」를 통해 가입한 회원이 플랫폼 상 임무를 완성 완료한 경우, 해당 임무 가치의 3%를「을」에게 추가 장려금으로 지급한다.')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('단체 임무완성 장려금 = 임무완성 건의 임무가치 × 3%')
r.bold = True; r.font.size = Pt(11); r.font.name = 'Consolas'
doc.add_paragraph()

heading(doc, '6.2 단체 임무발행 장려금 ( Publishers奖励 / 任务发布奖励)', level=2)
body(doc, '「을」가 직접 플랫폼에 임무를 게시·발행한 경우, 임무 게시 시 설정한 임무가치의 10%를「을」에게 추가 장려금으로 지급한다.')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('단체 임무발행 장려금 = 임무가치 설정금액 × 10%')
r.bold = True; r.font.size = Pt(11); r.font.name = 'Consolas'
doc.add_paragraph()

heading(doc, '6.3 지급 시기 / 发放时间', level=2)
body(doc, '상기 장려금은 해당 임무완성 또는 게시 확인 후 30일 이내，「을」의 HRT 지갑으로 전송한다.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 7 – HRT Value & Exchange
# ═══════════════════════════════════════════════════
heading(doc, '제7조 (HRT 가치 및 교환 조건) / 第7条（HRT价值及兑换条件）', level=1)

heading(doc, '7.1 현 시점 단가 미부여 / 现阶段单价不赋予', level=2)
body(doc, 'HRT의 현금 환산 단가는 현재 시점에서 부여하지 않으며, 플랫폼 이용 활성화에 따라 향후 시장 가격이 형성될 것으로 예측한다.')
doc.add_paragraph()

heading(doc, '7.2 향후 교환 가능 조건 / 未来兑换条件', level=2)
body(doc, '다음 조건이 충족된 시점부터 HRT의 현금 교환이 가능하다.')
body(doc, '① 회원충족: 플랫폼 유효회원이 일정 규모 이상 도달한 경우', indent=0.5)
body(doc, '② 충전체계 완비: 회원이 HRT 충전하여 서비스/상품에 이용할 수 있는 체계 구축 완료', indent=0.5)
body(doc, '③流動성 확보: 토큰의 시장流动性 확보 (거래소 상장 또는 DEX 유동성 확보)', indent=0.5)
body(doc, '④ 사전 통보: 상기 조건 충족 시「갑」은 사전 서면통보完了条件属实的情况下优先续约权利有义务遵守平台规则的情况下同等条件属实', indent=0.5)
doc.add_paragraph()

heading(doc, '7.3 교환 방식 / 兑换方式', level=2)
body(doc, '교환 가능 시「을」는「갑」이 지정한 방법에 따라 HRT를 법정통화로 교환 신청할 수 있으며,「갑」은 교환 신청 접수 후 30영업일 이내 처리한다. 교환 시 발생하는 수수료는「을」가 부담한다.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 8 – Obligations
# ═══════════════════════════════════════════════════
heading(doc, '제8조 (단체 의무) / 第8条（团体义务）', level=1)
body(doc, '「을」은 다음과 같은 의무를진다.')
body(doc, '① 플랫폼에 가입시키는 모든 회원이 본인의 실명·실제 정보로 가입하도록 안내한다.', indent=0.5)
body(doc, '② 회원이 유효회원 조건(6개월 내 1회 APP 접속)을 충족할 수 있도록 정기 안내 및宣传活动을 실시한다.', indent=0.5)
body(doc, '③ 회원의 개인정보 보호 및 플랫폼 이용약관을 준수하도록 교육·관리한다.', indent=0.5)
body(doc, '④ 플랫폼에 임무 게시 시 위조 정보나 불법 내용이 포함되지 않도록 확인한다.', indent=0.5)
body(doc, '⑤「갑」의 요청 시, 가입자 수 및 유효회원 현황에 대한 월간 보고서를 제출한다.', indent=0.5)
doc.add_paragraph()

heading(doc, '제9조 (갑의 의무) / 第9条（甲方义务）', level=1)
body(doc, '「갑」은 다음과 같은 의무를진다.')
body(doc, '① 플랫폼의 안정적 운영 및 기술 유지·보수를 제공한다.', indent=0.5)
body(doc, '② HRT 지급 내역을「을」에게 정확히 통보한다.', indent=0.5)
body(doc, '③ 유효회원 검증 기준 및 절차를 투명하게 공개한다.', indent=0.5)
body(doc, '④ 교환 가능 조건 충족 시, 정당한 교환 신청을 합리적 기간 내에 처리한다.', indent=0.5)
body(doc, '⑤ 계약 기간 중「을」의 파트너 지위를 희석시키거나 불이익을 주는 행위를 하지 않는다.', indent=0.5)
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 10 – Term
# ═══════════════════════════════════════════════════
heading(doc, '제10조 (계약 기간 및 갱신) / 第10条（合同期限及续期）', level=1)
body(doc, '① 本合同期限为 3년(3年)으로, 계약일로부터 효력이 발생한다.\n② 任一方可在合同到期前 60일(60天)书面通知对方续期或终止。\n③ 갱신 시 조건은 상호 협의를 통해 조정할 수 있다.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 11 – Termination
# ═══════════════════════════════════════════════════
heading(doc, '제11조 (계약 해지) / 第11条（合同解除）', level=1)
body(doc, '① 任一方严重违约且在对方书面催告后 30일(30天)内未纠正的，另一方可书面解除本合同。\n② 플랫폼 영구 운영 중단 시,「갑」은「을」에게 남은 HRT의 현금 환불을 보장한다.\n③ 계약 해지 시「을」에게 이미 부여된 HRT는「을」의 소유로 유지된다.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 12 – Liability
# ═══════════════════════════════════════════════════
heading(doc, '제12조 (손해배상 및 면책) / 第12条（损害赔偿及免责）', level=1)
body(doc, '① 불가항력(천재지변, 전쟁, 전염병, 정부 규제 등)으로 플랫폼 운영이 중단된 경우,「갑」은 면책된다.\n②「을」가 회원의 위법행위나 개인정보 유출 등 본 계약 관련 의무 위반으로「갑」 또는 제3자에게 손해를 끼힌 경우,「을」는 이를 배상한다.\n③ HRT 가격 변동으로 인한 투자 손실은「을」의 책임이며「갑」은 면책된다.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 13 – Confidentiality
# ═══════════════════════════════════════════════════
heading(doc, '제13조 (기밀 유지) / 第13条（保密条款）', level=1)
body(doc, '양측은 계약 이행 과정에서 취득한对方的商业情報, 技术信息, 会员个人信息를 제3자에게 공개하지 않는다. 本合同终止后仍持續有效。')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 14 – Dispute Resolution
# ═══════════════════════════════════════════════════
heading(doc, '제14조 (재판 관할 및 준거법) / 第14条（争议解决及适用法律）', level=1)
body(doc, '① 本계약는 중화인민공화국 법률에 따라 해석·준거된다.\n② 本合同引起的任何争议，双方应首先通过友好协商解决；协商不成的，任一方可向甲方所在地人民法院(深圳市中级人民法院)에 소송을 제기한다.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  ARTICLE 15 – Miscellaneous
# ═══════════════════════════════════════════════════
heading(doc, '제15조 (기타) / 第15条（其他）', level=1)
body(doc, '① 本合同自双方签字（盖章）之日起生效。\n② 本合同的任何修改或补充须经双方书面协议方可生效。\n③ 本合同一式 2부(两份)로, 각 국가 언어(한국어·중국어) 병기하며, 양측 각 1부씩 보관한다.\n   양문본의內容如有抵触，以中文本为准。\n④ 본 계약서에 기재되지 않은 사항에 대해서는 상호 협의로 처리한다.')
doc.add_paragraph()

# ═══════════════════════════════════════════════════
#  SIGNATURE BLOCK
# ═══════════════════════════════════════════════════
doc.add_page_break()
heading(doc, '서명란 / 签署栏', level=1, center=True)
doc.add_paragraph()

sig = doc.add_table(rows=6, cols=3)
sig.style = 'Table Grid'
sig.alignment = WD_TABLE_ALIGNMENT.CENTER

sig_headers = ['', '갑 (甲方)', '을 (乙方)']
for i, h in enumerate(sig_headers):
    sig.rows[0].cells[i].text = h
    for p in sig.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(10)
    add_shaded_table_row(sig, 0, 'D9E1F2')

sig_data = [
    ('상호 / 名称', '____________', '____________'),
    ('대표자 / 法定代表人', '____________', '____________'),
    ('직위 / 职务', '____________', '____________'),
    ('서명 / 签字', '__________________', '__________________'),
    ('날짜 / 日期', '2026년 ___월 ___일', '2026년 ___월 ___일'),
]
for ri, (a, b, c) in enumerate(sig_data):
    sig.rows[ri+1].cells[0].text = a
    sig.rows[ri+1].cells[1].text = b
    sig.rows[ri+1].cells[2].text = c
    for cell in sig.rows[ri+1].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()

# ─── Appendix ───
doc.add_page_break()
heading(doc, '부록: HRT 분배 계산 예시 (참고용) / 附录：HRT分配计算示例（仅供参考）', level=1)
body(doc, '아래 수치는 예시이며, 실제 수치는 플랫폼 런칭 후 확정된다. / 以下数值为示例，实际数以平台上线后确定为准。')
doc.add_paragraph()

add_table(doc,
    ['항목 / 项目', '예시 수치 / 示例数值'],
    [
        ['HRT 총 발행량', '1,000,000 HRT'],
        ['개발자 보유 (30%)', '300,000 HRT'],
        ['단체 풀 (50%)', '500,000 HRT'],
        ['개인 풀 (20%)', '200,000 HRT'],
        ['전체 파트너 유효회원 합계', '10,000명 (가정)'],
        ['1인당 HRT 단가', '50 HRT / 名 (역산치)'],
    ]
)
doc.add_paragraph()

heading(doc, '단체 A 수령 예상액 (예시)', level=2)
add_table(doc,
    ['항목', '산식', '예상 HRT'],
    [
        ['기본 HRT', '500명 × 50 HRT', '25,000 HRT'],
        ['임무완성 장려금(3%)', '활동 규모에 따라', '추가'],
        ['임무발행 장려금(10%)', '게시 임무 수에 따라', '추가'],
    ]
)
doc.add_paragraph()
body(doc, '* 본 계약서는 참고용 초안이며, 실제 체결 전 반드시 법률 자문을 받을 것을 권장한다. / 本合同为参考用初稿，实际签约前建议咨询专业律师。')

# ─── Save ───
out = r'E:\WorkBuddy\heartchain\docs\HeartChain_파트너십계약서_초안.docx'
doc.save(out)
print(f'Saved: {out}')
